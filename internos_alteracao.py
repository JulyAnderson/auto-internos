import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ALIGN_VERTICAL


# Import 'dados.csv' and transform it into a DataFrame
alteracao = pd.read_csv('alteracao.csv', usecols=["Mat",'Dia da Alteração',
                                               'Entrada 01','Saída Intervalo 01',
                                               'Entrada Intervalo 01','Saída 01',
                                               'Entrada 02','Saída 02',
                                               'FEZ CARGA SUPLEMENTAR NESTE DIA?'], sep=',')
alteracao.fillna('--', inplace=True)

# Import 'base.csv' and transform it into a DataFrame
homologado = pd.read_csv('homologado.csv', usecols=["Mat", "Nome", "Sala 2ª à 6ª feira",
                                               "PPM 2ª feira", "HTPC 3ª feira", "AP", "HL", "HLE",
                                               "Horas Compensadas", "C.H. Semanal"], sep=",")
homologado.fillna('--', inplace=True)

# Filtrar matrículas das alterações em que não houve cargas suplementares
matriculas = alteracao.loc[(alteracao['FEZ CARGA SUPLEMENTAR NESTE DIA?'] != 'Sim') &
                           (alteracao['Mat'].isin(homologado['Mat'])), 'Mat'].unique()

print("""
____________________________________________________________________________________
Bem-vindo ao gerador de Internos de Alteração de Horário Docentes da Escola Modesto.
------------------------------------------------------------------------------------

Antes de seguir com o programa, confira as etapas:

1- Download da planilha de horários alterados, do GoogleDrive, no formato csv, o arquivo deve se chamar 'alteracao';
2- Download da planilha de horários homologados, do GoogleDrive, no formato csv, o arquivo deve se chamar 'homologado';
3- O programa deve estar na mesma pasta dos arquivos baixados anteriormente;
4- Nesta pasta deve haver uma outra pasta chamada internos, que receberá os internos novos.


""")


mes_pagamento = input("Informe a qual mês esse pagamento é referente. Exemplo: junho-2023= ")
numero = int(input("Digite o número do Primeiro Interno. Exemplo: 93= "))
data_local = input('Digite a data de envio do Interno. Exemplo: 01 de junho de 2023= ')
responsavel_assinatura = input("""Quem é o responsável pela assinatura dos internos? Selecione uma das opções:
1- Silvia Elaine
2- Anelisa Luciene
3- Outros: """)

if responsavel_assinatura == "1":
    nome_responsavel = "_____________________________\nSilvia Elaine da Silva Ganzela\nDiretora de Escola"
elif responsavel_assinatura == "2":
    nome_responsavel = "_____________________________\nAnelisa Luciene Lopes Regovich\nCoordenadora Pedagógica"
else:
    nome_responsavel = "_____________________________\nResponsável"

caminho_pasta = 'internos/'

for matricula in matriculas:
    homologado_filtrado = homologado[homologado['Mat'] == matricula]

    interno = Document('modelo.docx')

    numero_interno = interno.add_paragraph()
    numero_interno.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    numero_interno_run = numero_interno.add_run(f'Interno nº {numero}/ 2023')
    numero_interno_run.bold = True
    numero_interno_run.font.size = Pt(12)
    interno.add_paragraph()
    interno.add_paragraph()

    data_interno = interno.add_paragraph()
    data_interno.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    data_interno_run = data_interno.add_run(f'São José do Rio Preto, {data_local}.')
    data_interno_run.font.size = Pt(12)
    interno.add_paragraph()
    interno.add_paragraph()

    destinatario_interno = interno.add_paragraph()
    destinatario_interno.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    destinatario_interno_run = destinatario_interno.add_run("""À Secretaria Municipal de Administração
A/C Coordenadoria de Pagamento 
Bancada da Educação""")
    destinatario_interno_run.bold = True
    destinatario_interno_run.font.size = Pt(12)
    interno.add_paragraph()
    interno.add_paragraph()

    assunto_interno = interno.add_paragraph()
    assunto_interno_run = assunto_interno.add_run("Assunto: Alteração pontual de horário.")
    assunto_interno_run.bold = True
    assunto_interno_run.font.size = Pt(12)
    interno.add_paragraph()
    interno.add_paragraph()

    corpo_interno = interno.add_paragraph()
    corpo_interno.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    corpo_interno_run = corpo_interno.add_run("""Encaminham-se as alterações pontuais de horários do servidor abaixo descrito:


Horário Homologado: """)
    corpo_interno_run.font.size = Pt(12)
    interno.add_paragraph()
    
    # Add the table for Homologated Schedule
    table_homologado = interno.add_table(rows=homologado_filtrado.shape[0] + 1, cols=homologado_filtrado.shape[1])
    table_homologado.style = 'Table Grid'  # Add borders to the table

    # Set alignment for each cell in the table
    for row in table_homologado.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set the header row cells' text and formatting
    header_row = table_homologado.rows[0]
    for i, column_name in enumerate(homologado_filtrado.columns):
        header_cell = header_row.cells[i]
        header_cell.text = column_name
        header_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header_cell.bold = True

    # Set the remaining cells' text and formatting
    for row in range(homologado_filtrado.shape[0]):
        for col in range(homologado_filtrado.shape[1]):
            cell = table_homologado.cell(row + 1, col)
            cell.text = str(homologado_filtrado.values[row, col])
            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    # Adiciona um paragráfo como separador 
    interno.add_paragraph()

    # Filtra as interações para a matricula atual
    alteracoes_filtradas = alteracao[alteracao['Mat'] == matricula]
    alteracoes_filtradas = alteracoes_filtradas[alteracoes_filtradas['FEZ CARGA SUPLEMENTAR NESTE DIA?'] != 'Sim']
    colunas_desejadas = ['Dia da Alteração',
                        'Entrada 01',
                        'Saída Intervalo 01',
                        'Entrada Intervalo 01',
                        'Saída 01',
                        'Entrada 02',
                        'Saída 02']
    alteracoes_filtradas = alteracoes_filtradas.loc[:, colunas_desejadas]

    horario_alterado_interno = interno.add_paragraph()
    horario_alterado_interno.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    horario_alterado_interno = horario_alterado_interno.add_run("Horário Alterado para: ")
    horario_alterado_interno.font.size = Pt(12)
    interno.add_paragraph()

    # Adiciona uma tabela para as alterações
    table_alteracoes = interno.add_table(rows=alteracoes_filtradas.shape[0] + 1, cols=alteracoes_filtradas.shape[1])
    table_alteracoes.style = 'Table Grid'  # Add borders to the table

    # Set alignment for each cell in the table
    for row in table_alteracoes.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set the header row cells' text and formatting
    header_row = table_alteracoes.rows[0]
    for i, column_name in enumerate(alteracoes_filtradas.columns):
        header_cell = header_row.cells[i]
        header_cell.text = column_name
        header_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header_cell.bold = True

    # Set the remaining cells' text and formatting
    for row in range(alteracoes_filtradas.shape[0]):
        for col in range(alteracoes_filtradas.shape[1]):
            cell = table_alteracoes.cell(row + 1, col)
            cell.text = str(alteracoes_filtradas.values[row, col])
            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add a paragraph for the responsible person
    interno.add_paragraph()
    interno.add_paragraph()
    responsavel_interno = interno.add_paragraph()
    responsavel_interno.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    responsavel_interno_run = responsavel_interno.add_run(nome_responsavel)
    responsavel_interno_run.bold = True
    responsavel_interno_run.font.size = Pt(12)


    # Salva o documento com nome único 
    interno.save(caminho_pasta + f'Interno nº{numero}-2023 - Alteração de horário de {homologado_filtrado.iloc[0]["Nome"]}-{mes_pagamento}.docx')

    numero += 1


print("""

PROCESSO ENCERRADO, TODOS OS INTERNOS FORAM GERADOS.

  #######  #######   ########
  #     #  #     #   #
  #     #  #  ###    #   ####
  #     #  #     #   #      #
  #######  ######    ########
                            #
""")